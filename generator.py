from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


# section generator
class SectionGenerator:
    def __init__(self):
        pass

    def insertHR(self, paragraph):
        p = paragraph._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
 
    def create_section(self, document, section_name):
        # section_header = document.add_heading(level=5)
        section_header = document.add_paragraph()

        runner = section_header.add_run(f'{section_name}')
        runner.bold = True
        runner.size = Pt(13)

        self.insertHR(section_header)


    def add_item(self, document, title, time, detail, bullet_point_array):
        EXACTLY=1.15
        title_para = document.add_paragraph()
        runner = title_para.add_run(title)
        runner.bold = True
        
        title_para.add_run(f'\t{time}')  # tab will trigger tabstop
        sec = document.sections[0]
        # finding end_point for the content 
        margin_end = Inches(
            sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
        tab_stops = title_para.paragraph_format.tab_stops
        # adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned.
        tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
        title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        title_para.paragraph_format.space_after = Pt(2)

        b=document.add_paragraph(f"{detail}")
        b.paragraph_format.space_after = Pt(2)
        b.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        for i in bullet_point_array:
            c=document.add_paragraph(i, style='List Bullet')
            c.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            c.paragraph_format.space_after = Pt(2)


        
    def create_skills(self, document, array):
        section_header = document.add_heading(level=5)
        runner = section_header.add_run('Skills')
        runner.bold = True
        self.insertHR(section_header)
        if len(array) == 1:
            document.add_paragraph(array)
        else:
            para = document.add_paragraph(f"{array[0]}\n")
            for i in range(1, len(array)):
                para.add_run(f"{array[i]}\n")

    def create_hobbies(self, document, array):
        section_header = document.add_heading(level=5)
        runner = section_header.add_run('Hobbies')
        runner.bold = True
        self.insertHR(section_header)
        if len(array) == 1:
            document.add_paragraph(array)
        else:
            para = document.add_paragraph(f"{array[0]}")
            for i in range(1, len(array)):
                para.add_run(f"{array[i]}")
        

### Writing The Resume ###

data = {"name": "Bryan Lim Kai Wen", "mobile": "+65 12345678", "email": "rando@ntu.edu.sg"}
document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

document.add_paragraph('')
# write resume header info
title= document.add_paragraph()
title_format = title.paragraph_format
title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_runner = title.add_run(f'{data["name"]} | Mobile No.: {data["mobile"]} | Email: {data["email"]}')
title_runner.bold = True
title_runner.size = Pt(10)

section_generator = SectionGenerator()

section_generator.create_section(document, "Education")
section_generator.add_item(document, "Nanyang Technological University", "hihidshif", "asoidjfjk ka;l dkjfl;aj flkja ljdlkasjfalka kfa;ls", ["123123","12312312",'sdfsdf'])


section_generator.create_skills(document, ["123123","12312312",'sdfsdf'])
section_generator.create_hobbies(document, ["123123","12312312",'sdfsdf'])


########################################################################################################################################################################
'''

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))


document.add_page_break()

'''
###########################################################################################################################################################################

document.save('demo.docx')